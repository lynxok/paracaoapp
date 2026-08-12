from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

OUT = Path(__file__).with_name('Manual_de_Usuario_Optica_Paracao.docx')
CAPTURAS = Path(__file__).with_name('capturas')

def font(run, size=None, bold=None, color=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    if size: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def shade(cell, value):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), value); tcPr.append(shd)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run('Óptica Paracao  |  Página ')
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.8)
sec.left_margin = sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.18
for name, size, color, before, after in [('Heading 1', 17, (30, 82, 133), 18, 8), ('Heading 2', 13, (30, 82, 133), 12, 6), ('Heading 3', 11.5, (31, 77, 120), 8, 4)]:
    s = styles[name]; s.font.name = 'Calibri'; s.font.size = Pt(size); s.font.color.rgb = RGBColor(*color); s.font.bold = True
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)

header = sec.header.paragraphs[0]
header.text = 'ÓPTICA PARACAO  ·  MANUAL DE USUARIO'
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: font(r, 8.5, True, (100, 116, 139))
add_page_number(sec.footer.paragraphs[0])

def p(text='', style=None, bold_start=None):
    par = doc.add_paragraph(style=style)
    if bold_start and text.startswith(bold_start):
        font(par.add_run(bold_start), bold=True)
        font(par.add_run(text[len(bold_start):]))
    else: font(par.add_run(text))
    return par

def bullets(items):
    for item in items:
        par = doc.add_paragraph(style='List Bullet')
        par.paragraph_format.space_after = Pt(3)
        font(par.add_run(item))

def steps(items):
    for item in items:
        par = doc.add_paragraph(style='List Number')
        par.paragraph_format.space_after = Pt(4)
        font(par.add_run(item))

def note(title, text):
    table = doc.add_table(rows=1, cols=1); table.autofit = False; table.columns[0].width = Inches(6.8)
    cell = table.cell(0, 0); shade(cell, 'EAF2F8')
    par = cell.paragraphs[0]; font(par.add_run(title + ' '), bold=True, color=(30,82,133)); font(par.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def screenshot(filename, caption):
    image_path = CAPTURAS / filename
    if not image_path.exists():
        raise FileNotFoundError(f'No se encontró la captura: {image_path}')
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.keep_with_next = True
    par.add_run().add_picture(str(image_path), width=Inches(6.75))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    font(cap.add_run(caption), 8.5, False, (100,116,139))

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(80)
cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(cover.add_run('MANUAL DE USUARIO'), 28, True, (30,82,133))
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(sub.add_run('ERP/CRM para la gestión integral de Óptica Paracao'), 15, False, (71,85,105))
doc.add_paragraph().paragraph_format.space_after = Pt(30)
intro = doc.add_paragraph(); intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(intro.add_run('Guía operativa para ventas, atención al cliente, pedidos ópticos, inventario, proveedores y finanzas.'), 11)
doc.add_paragraph().paragraph_format.space_after = Pt(105)
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(meta.add_run('Versión del sistema revisada: agosto de 2026'), 9.5, False, (100,116,139))
doc.add_page_break()

doc.add_heading('Cómo usar esta guía', level=1)
p('Este manual describe el uso cotidiano del sistema. Los nombres de campos, botones y pestañas pueden variar ligeramente según los permisos asignados a cada usuario.')
note('Recomendación:', 'antes de operar, confirme que está en la sucursal correcta y que la caja seleccionada es la que corresponde a la operación.')
doc.add_heading('Índice', level=2)
bullets(['1. Ingreso y navegación general', '2. Panel de control', '3. Clientes y cuenta corriente', '4. Ventas, pedidos y recetas ópticas', '5. Proveedores, compras y facturas', '6. Inventario y movimientos de stock', '7. Caja, finanzas y conciliación', '8. Reportes, marketing, configuración y soporte'])

doc.add_heading('1. Ingreso y navegación general', level=1)
p('Óptica Paracao es una aplicación web y no requiere instalar programas, Node.js ni dependencias. Abra un navegador actualizado e ingrese directamente a https://opticagestionparacao.lnx.com.ar, o utilice el enlace web entregado por el administrador. Para acceder necesita conexión a Internet y sus credenciales. Luego utilice el menú lateral para acceder a cada módulo.')
screenshot('01-acceso.png', 'Pantalla de acceso, capturada sin credenciales visibles.')
doc.add_heading('Reglas de operación', level=2)
bullets(['Verifique el cliente antes de crear una venta o pedido.', 'Registre cada cobro en la caja o cuenta de destino correcta.', 'No complete un pedido hasta confirmar la entrega y el saldo.', 'Registre ingresos y egresos de stock para mantener las existencias actualizadas.'])

doc.add_heading('2. Panel de control', level=1)
p('El Panel de Control resume la operación diaria. Allí se visualizan métricas de ventas, pedidos en taller, saldos pendientes y accesos rápidos.')
screenshot('02-inicio.png', 'Panel principal con información identificatoria sanitizada.')
doc.add_heading('Acciones frecuentes', level=2)
bullets(['Abrir una venta rápida para productos sin receta.', 'Crear un pedido óptico para una receta nueva.', 'Ir a clientes para altas, cobros o consulta de historial.', 'Acceder a caja para registrar movimientos o hacer el cierre.'])

doc.add_heading('3. Clientes y cuenta corriente', level=1)
p('El módulo Clientes reúne la ficha personal, cobertura médica, historial de pedidos y situación financiera de cada persona.')
screenshot('03-clientes.png', 'Gestión de Clientes; los registros fueron ocultados para proteger datos personales.')
doc.add_heading('Dar de alta o editar un cliente', level=2)
steps(['Abra Clientes y use la búsqueda para comprobar si ya existe un registro por nombre, DNI o teléfono.', 'Seleccione Nuevo cliente o el icono de edición.', 'Complete los datos personales y de contacto. Registre obra social o cobertura cuando corresponda.', 'Guarde los cambios y confirme que la ficha aparezca en el listado.'])
doc.add_heading('Consultar cuenta corriente y pedidos', level=2)
steps(['Desde la fila del cliente, abra Cuenta Corriente para revisar cargos, pagos, señas y saldo.', 'Para registrar un pago, indique importe, medio de pago y caja o cuenta de destino.', 'Abra Ver pedidos para consultar tipo de trabajo, estado, total y saldo pendiente.', 'Emita el comprobante si el flujo de trabajo lo requiere.'])
note('Importante:', 'un saldo pendiente debe mantenerse asociado al cliente hasta registrar el cobro efectivo.')

doc.add_heading('4. Ventas, pedidos y recetas ópticas', level=1)
doc.add_heading('Venta rápida (no recetados)', level=2)
screenshot('04-ventas-rapidas.png', 'Pantalla de Ventas Rápidas.')
steps(['Abra Venta No Recetados.', 'Busque y agregue productos de stock al carrito.', 'Revise cantidades, precio y descuentos autorizados.', 'Seleccione el cliente si corresponde y confirme el medio de pago.', 'Finalice la venta: el sistema registra el ingreso e impacta el stock.'])
doc.add_heading('Crear un pedido óptico', level=2)
screenshot('05-pedidos.png', 'Selección del tipo de pedido recetado.')
screenshot('06-pedido-monofocal.png', 'Formulario real de Detalle de Pedido Monofocal, vacío y sanitizado.')
steps(['Abra Nuevo Pedido y seleccione el tipo de solución: monofocal, multifocal o lente de contacto.', 'Seleccione o cree el cliente.', 'Cargue los valores de receta solicitados: OD/OI, esférico, cilíndrico, eje, adición, distancia pupilar y, si aplica, altura de montaje.', 'Elija armazón, cristales y tratamientos. Verifique disponibilidad de stock por sucursal.', 'Defina el total, seña o pago completo y el destino del cobro.', 'Confirme el pedido. El estado inicial normalmente será En Taller.'])
doc.add_heading('Seguimiento y entrega', level=2)
bullets(['En Taller: trabajo enviado o pendiente de procesamiento.', 'Para Retirar: trabajo terminado y disponible para el cliente.', 'Completado: pedido entregado y cerrado. Antes de marcarlo, registre cualquier saldo restante.'])

doc.add_heading('5. Proveedores, compras y facturas', level=1)
screenshot('08-proveedores.png', 'Pantalla de Proveedores; los registros fueron ocultados.')
doc.add_heading('Registrar proveedor', level=2)
p('En Proveedores, cree o edite la ficha con razón social, CUIT, datos de contacto, rubro, datos bancarios y condición de pago habitual.')
doc.add_heading('Cargar una compra', level=2)
steps(['Abra Proveedores y seleccione la sección de Compras.', 'Elija Cargar nueva compra.', 'Indique proveedor, tipo y número de comprobante, fechas de emisión y vencimiento, importe y condición de pago.', 'Agregue una descripción que permita identificar la operación.', 'Confirme. La compra impactará en la cuenta corriente del proveedor y, si corresponde, en el inventario.'])
doc.add_heading('Controlar facturas pendientes', level=2)
bullets(['Abra Facturas Pendientes para ordenar obligaciones por vencimiento.', 'Revise importes, fechas y estado para priorizar pagos.', 'Al pagar, genere la orden de pago y seleccione la caja o banco de donde sale el dinero.', 'Compruebe que el saldo del proveedor disminuya después de confirmar.'])
doc.add_heading('Laboratorios', level=2)
screenshot('09-laboratorios.png', 'Pantalla de Laboratorios; los registros fueron ocultados.')
p('Use esta sección para controlar trabajos derivados y su liquidación. Actualice un estado solamente cuando exista respaldo del taller o laboratorio.')

doc.add_heading('6. Inventario y movimientos de stock', level=1)
p('El inventario administra productos por SKU, categoría y sucursal. Incluye armazones, cristales, lentes de contacto, lentes de sol y accesorios.')
screenshot('07-stock.png', 'Pantalla de Stock; los registros fueron ocultados.')
doc.add_heading('Alta y ajuste de productos', level=2)
steps(['Abra Gestión de Stock y seleccione Nuevo producto o edite uno existente.', 'Complete SKU, nombre, categoría, precios y umbral de stock crítico.', 'Registre el stock por sucursal cuando corresponda y guarde.'])
doc.add_heading('Movimientos', level=2)
bullets(['Ingreso: utilícelo al recibir mercadería; asocie proveedor y factura si están disponibles.', 'Egreso: registre ajustes, roturas, devoluciones u otras salidas justificadas.', 'Mover stock: elija sucursal de origen, destino y cantidad. El sistema valida que haya unidades disponibles.', 'Revise Movimientos de Stock para auditar el historial.'])
note('Alerta de stock:', 'cuando una existencia alcanza el mínimo configurado, planifique la reposición antes de confirmar nuevas ventas.')

doc.add_heading('7. Caja, finanzas y conciliación', level=1)
screenshot('10-caja-finanzas.png', 'Caja y Finanzas; movimientos e identidad fueron sanitizados.')
doc.add_heading('Registrar movimientos de caja', level=2)
steps(['Abra Caja y Finanzas y seleccione la caja o cuenta correspondiente.', 'Use Nuevo movimiento para cargar ingresos o egresos no registrados desde una venta o pago.', 'Indique concepto, categoría, importe, fecha y comprobante de respaldo cuando exista.', 'Confirme y revise el saldo actualizado.'])
doc.add_heading('Cierre de caja', level=2)
steps(['Al finalizar la jornada, abra Cierre de Caja.', 'Cuente el efectivo físico y cargue el monto real.', 'Compare el resultado con el saldo teórico. Si existe diferencia, revise los movimientos antes de confirmar.', 'Registre observaciones y confirme el cierre cuando esté validado.'])
doc.add_heading('Conciliación bancaria', level=2)
steps(['Abra la pestaña Conciliación.', 'Seleccione banco o cuenta y defina el rango de fechas según el extracto.', 'Ingrese el saldo final informado por el banco y, de existir, los gastos bancarios.', 'Compare los movimientos mostrados con el extracto y márquelos como conciliados.', 'Cuando la diferencia sea cero y las transacciones estén revisadas, confirme la conciliación.'])

doc.add_heading('8. Reportes, marketing, configuración y soporte', level=1)
doc.add_heading('Borradores de facturación', level=2)
screenshot('11-borradores-facturacion.png', 'Borradores de Facturación; los registros fueron ocultados.')
p('Revise cliente, condición fiscal, detalle, importes, tipo de comprobante y punto de venta antes de procesar un borrador. La emisión fiscal debe realizarse únicamente con autorización.')
doc.add_heading('Reportes', level=2)
screenshot('12-reportes.png', 'Pantalla de Reportes.')
p('Reportes permite analizar ventas, egresos, recetas y rentabilidad. Aplique los filtros disponibles antes de consultar o exportar resultados.')
doc.add_heading('CRM y marketing', level=2)
screenshot('13-crm-marketing.png', 'CRM y Marketing; datos identificatorios sanitizados.')
p('El módulo CRM & Marketing ayuda a segmentar clientes y gestionar acciones de fidelización o recordatorios. Confirme siempre el segmento, el contenido y los permisos antes de realizar un envío.')
doc.add_heading('Configuración', level=2)
screenshot('14-ajustes.png', 'Pantalla de Ajustes.')
bullets(['Administre sucursales, usuarios y permisos únicamente con una cuenta autorizada.', 'Actualice categorías, medios de pago y parámetros operativos antes de que se utilicen en nuevas transacciones.', 'Mantenga las copias de seguridad y la información de acceso bajo control del administrador.'])
doc.add_heading('Ayuda y soporte', level=2)
screenshot('15-ayuda.png', 'Pantalla de Ayuda y manual guiado.')
p('Utilice Ayuda y Soporte para consultar preguntas frecuentes y flujos del sistema. Cuando informe una incidencia, indique el módulo, la acción realizada, el mensaje visible y una captura si es posible.')

doc.add_heading('9. Complemento operativo basado en el plan QA', level=1)
doc.add_heading('Alta, búsqueda e historial de clientes', level=2)
screenshot('16-alta-cliente.png', 'Formulario vacío de alta de cliente.')
p('Busque primero por DNI o nombre para evitar duplicados. Complete nombre, apellido, identificación, nacimiento, contacto, dirección, obra social y afiliado. Si aparece una advertencia de DNI duplicado, cancele y abra la ficha existente. Use el historial para comparar pedidos y recetas; preserve los registros asociados.')

doc.add_heading('Inventario: productos e ingresos', level=2)
screenshot('17-alta-producto.png', 'Formulario vacío de alta de producto.')
screenshot('18-ingreso-mercaderia.png', 'Formulario vacío de ingreso de mercadería.')
bullets(['Use un SKU único y precios no negativos.', 'Registre stock por sucursal y umbral crítico.', 'Respalde ingresos con proveedor, remito o factura.', 'Para egresos o transferencias, indique motivo, origen, destino y cantidad; nunca transfiera más stock que el disponible.'])

doc.add_heading('Recetas multifocales y lentes de contacto', level=2)
screenshot('19-pedido-multifocal.png', 'Formulario real de Multifocales/Bifocales.')
p('En multifocales, revise ESF, CIL y EJE de lejos, adiciones de OD/OI, DNP de lejos y cerca y altura. Confirme que el cristal admita el rango de la receta.')
screenshot('20-lentes-contacto.png', 'Formulario real de Lentes de Contacto.')
p('En lentes de contacto cargue por ojo ESF, curvatura BC, diámetro DIA y color, además de marca y frecuencia cuando corresponda. No sustituya medidas faltantes por valores típicos.')

doc.add_heading('Caja: egresos, conciliación y anulaciones', level=2)
screenshot('21-egreso-caja.png', 'Formulario de registro de egreso.')
p('Seleccione concepto, caja, monto, fecha, categoría y medio de pago. Use Orden de Pago para proveedores cuando corresponda. No borre movimientos para corregirlos: aplique el procedimiento autorizado de anulación.')
screenshot('22-conciliacion.png', 'Pantalla de conciliación bancaria.')
p('Compare cada movimiento con el extracto, registre gastos documentados y cierre únicamente cuando la diferencia esté explicada o sea cero.')

doc.add_heading('Configuraciones administrativas', level=2)
screenshot('23-ajustes-usuarios.png', 'Gestión de usuarios; filas sanitizadas.')
screenshot('24-ajustes-obras-sociales.png', 'Configuración de obras sociales.')
screenshot('25-ajustes-bancos.png', 'Configuración de bancos y entidades.')
screenshot('26-ajustes-cristales.png', 'Catálogo y rangos de cristales; filas sanitizadas.')
screenshot('27-audit-log.png', 'Registro de auditoría; eventos sanitizados.')
bullets(['Otorgue a cada usuario solo los permisos necesarios.', 'Revise coberturas antes de aplicarlas a pedidos.', 'Mantenga bancos y cajas digitales asociados a la sucursal correcta.', 'No amplíe rangos de cristales sin confirmación técnica.', 'Use Audit Log únicamente para consulta y trazabilidad.'])

doc.add_heading('Seguridad y validaciones', level=2)
bullets(['Cierre sesión desde el perfil y no comparta cuentas.', 'No ejecute casos QA ni cargue datos ficticios en producción.', 'No ignore alertas de DNI duplicado, stock insuficiente o receta fuera de rango.', 'Oculte datos clínicos, personales, bancarios y fiscales en capturas.', 'En pantallas pequeñas, confirme que puede ver todo el formulario y el botón de acción antes de cargar.'])

doc.add_heading('Lista diaria de verificación', level=1)
bullets(['Confirmar caja y sucursal antes de comenzar.', 'Registrar ventas, señas, cobros y egresos en el momento.', 'Actualizar los estados de pedidos en taller y para retirar.', 'Controlar alertas de stock y facturas próximas a vencer.', 'Realizar el cierre de caja y revisar diferencias al finalizar la jornada.'])

doc.save(OUT)
print(OUT)
